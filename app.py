"""
DataSentry v2 — Gradio UI for Hugging Face Spaces
Runs the full 4-layer detection pipeline inline — no separate API needed.
"""

import os
import csv
import html
import json
import tempfile
import gradio as gr
from pathlib import Path
from dotenv import load_dotenv
load_dotenv()

from src.detector import DataSentryDetector

detector = DataSentryDetector(
    spacy_model="en_core_web_sm",
    audit_db="/tmp/datasentry_audit.db"
)

SAMPLES = {
    "Patient record": (
        "Patient John Smith, DOB: 03/15/1978, SSN 432-56-7890 was admitted to "
        "City Medical Center. MRN: MRN-2024-00142. Diagnosis: E11.9 (Type 2 Diabetes). "
        "Prescribed metformin 500mg. Contact: jsmith@email.com | (555) 867-5309"
    ),
    "Financial data": (
        "Account holder: Sarah Johnson, account number 8834291055, "
        "routing 021000021. Credit card: 4532015112830366. "
        "Address: 742 Evergreen Terrace, Springfield, IL 62701."
    ),
    "Ambiguous text": (
        "The patient was seen at the clinic in January. "
        "Her doctor noted concerning values. She takes daily medication. "
        "Insurance member id: MBR-2024-99182."
    ),
}

CHAR_LIMIT = 50_000


# ── File ingestion ─────────────────────────────────────────────────────────────

def _extract_pdf(path):
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            return "\n\n".join(p.extract_text() or "" for p in pdf.pages)
    except Exception:
        try:
            from pypdf import PdfReader
            return "\n\n".join(p.extract_text() or "" for p in PdfReader(path).pages)
        except Exception as e:
            return f"[PDF error: {e}]"

def _extract_csv(path):
    rows = []
    with open(path, newline="", encoding="utf-8-sig") as f:
        for row in csv.reader(f):
            rows.append(", ".join(row))
    return "\n".join(rows)

def _extract_docx(path):
    try:
        import docx
        doc = docx.Document(path)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts)
    except Exception as e:
        return f"[DOCX error: {e}]"

def _extract_xlsx(path):
    try:
        from openpyxl import load_workbook
        wb = load_workbook(path, read_only=True, data_only=True)
        parts = []
        for sheet in wb.sheetnames:
            ws = wb[sheet]
            parts.append(f"### Sheet: {sheet}")
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) for c in row if c is not None and str(c).strip()]
                if cells:
                    parts.append(", ".join(cells))
        wb.close()
        return "\n".join(parts)
    except Exception as e:
        return f"[XLSX error: {e}]"

def ingest_file(file_path):
    if not file_path:
        return "", ""
    ext = Path(file_path).suffix.lower()
    try:
        if ext == ".pdf":
            text = _extract_pdf(file_path)
        elif ext == ".csv":
            text = _extract_csv(file_path)
        elif ext in (".docx", ".doc"):
            text = _extract_docx(file_path)
        elif ext in (".xlsx", ".xls"):
            text = _extract_xlsx(file_path)
        elif ext == ".txt":
            text = open(file_path, encoding="utf-8", errors="replace").read()
        else:
            return "", f"⚠ Unsupported file type: {ext}"
        return text, f"✓ Loaded {ext.upper()[1:]} — {len(text):,} chars extracted"
    except Exception as e:
        return "", f"✗ File error: {e}"


# ── Redaction ──────────────────────────────────────────────────────────────────

def build_redacted_text(source_text, entities):
    if not entities:
        return source_text
    sorted_ents = sorted(entities, key=lambda e: e.start, reverse=True)
    redacted = source_text
    for ent in sorted_ents:
        redacted = redacted[:ent.start] + f"[REDACTED {ent.entity_type}]" + redacted[ent.end:]
    return redacted

def build_redacted_file(original_path, source_text, entities):
    ext = Path(original_path).suffix.lower() if original_path else ".txt"
    out_ext = ext if ext in (".csv", ".txt") else ".txt"
    tmp = tempfile.NamedTemporaryFile(
        mode="w", suffix=out_ext, delete=False,
        encoding="utf-8", prefix="datasentry_redacted_"
    )
    tmp.write(build_redacted_text(source_text, entities))
    tmp.close()
    return tmp.name


# ── HTML rendering helpers ─────────────────────────────────────────────────────

def _conf_color(conf):
    if conf >= 0.85: return "#10b981"  # green
    if conf >= 0.70: return "#f59e0b"  # amber
    return "#ef4444"                    # red

def render_summary_html(result, file_status):
    file_banner = ""
    if file_status:
        file_banner = (
            f'<div class="ds-banner">{html.escape(file_status)}</div>'
        )

    budget_banner = ""
    if getattr(result, "claude_skipped_budget", False):
        budget_banner = (
            '<div class="ds-banner ds-banner-warn">'
            '⚡ <strong>Claude arbitration paused</strong> — daily demo budget '
            f'reached ({result.claude_calls_today}/{result.claude_budget_daily} '
            'calls). Low-confidence candidates were dropped; high-confidence '
            'regex matches still shown. Resets at 00:00 UTC.'
            '</div>'
        )

    layers = " · ".join(html.escape(l) for l in result.layers_used)
    budget_footer = ""
    if getattr(result, "claude_budget_daily", 0):
        calls   = getattr(result, "claude_calls_today", 0)
        budget  = result.claude_budget_daily
        budget_footer = (
            f' &nbsp;·&nbsp; Claude budget: <strong>{calls}/{budget}</strong> today'
        )
    return f"""
{file_banner}
{budget_banner}
<div class="ds-tiles">
  <div class="ds-tile pii">
    <div class="ds-tile-num">{result.total_pii}</div>
    <div class="ds-tile-lbl">PII detected</div>
  </div>
  <div class="ds-tile phi">
    <div class="ds-tile-num">{result.total_phi}</div>
    <div class="ds-tile-lbl">PHI detected</div>
  </div>
  <div class="ds-tile">
    <div class="ds-tile-num">{result.processing_ms:.0f}<span class="ds-unit">ms</span></div>
    <div class="ds-tile-lbl">Processing time</div>
  </div>
  <div class="ds-tile">
    <div class="ds-tile-num">{len(result.entities)}</div>
    <div class="ds-tile-lbl">Total entities</div>
  </div>
</div>
<div class="ds-meta">Layers used: <strong>{layers}</strong> &nbsp;·&nbsp; Run ID: <code>{result.run_id[:8]}</code>{budget_footer}</div>
"""

def render_entities_html(entities):
    if not entities:
        return (
            '<div class="ds-empty">'
            '<div class="ds-empty-ico">✓</div>'
            '<div class="ds-empty-msg">No PII or PHI detected</div>'
            '<div class="ds-empty-sub">This text appears safe to share.</div>'
            '</div>'
        )
    cards = []
    for e in entities:
        cat_class = e.category.lower()
        conf_pct  = int(round(e.confidence * 100))
        conf_col  = _conf_color(e.confidence)
        claude_badge = ''
        if e.claude_override or e.detection_layer == "claude":
            claude_badge = '<span class="ds-badge ds-badge-claude">Claude reviewed</span>'
        cards.append(f"""
<div class="ds-card ds-card-{cat_class}">
  <div class="ds-card-head">
    <div class="ds-card-type">
      <span class="ds-cat ds-cat-{cat_class}">{html.escape(e.category)}</span>
      <span class="ds-type">{html.escape(e.entity_type)}</span>
    </div>
    <div class="ds-card-conf">
      <span class="ds-conf-pct" style="color:{conf_col}">{conf_pct}%</span>
      <div class="ds-conf-bar"><div class="ds-conf-bar-fill" style="width:{conf_pct}%;background:{conf_col}"></div></div>
    </div>
  </div>
  <div class="ds-card-text">{html.escape(e.text[:80])}</div>
  <div class="ds-card-foot">
    <span class="ds-layer">via <strong>{html.escape(e.detection_layer)}</strong></span>
    {claude_badge}
  </div>
</div>
""")
    return f'<div class="ds-cards">{"".join(cards)}</div>'

def render_annotated_html(text, entities):
    if not entities:
        return f'<div class="ds-annot"><pre>{html.escape(text)}</pre></div>'
    parts = []
    cursor = 0
    for e in sorted(entities, key=lambda x: x.start):
        if e.start < cursor:  # skip overlap (shouldn't happen post-merge)
            continue
        parts.append(html.escape(text[cursor:e.start]))
        cat = e.category.lower()
        parts.append(
            f'<mark class="ds-mark ds-mark-{cat}" '
            f'title="{html.escape(e.entity_type)} · {int(e.confidence*100)}% via {html.escape(e.detection_layer)}">'
            f'{html.escape(text[e.start:e.end])}'
            f'<sup class="ds-mark-tag">{html.escape(e.entity_type)}</sup>'
            f'</mark>'
        )
        cursor = e.end
    parts.append(html.escape(text[cursor:]))
    return f'<div class="ds-annot"><pre>{"".join(parts)}</pre></div>'


# ── Entity label helper (for feedback checkboxes) ──────────────────────────────

def _entity_label(ent):
    return f"{ent.entity_type}  →  {ent.text[:40]}  ({ent.confidence:.0%} · {ent.detection_layer})"


# ── Main detection handler ─────────────────────────────────────────────────────

def detect(text, file_upload, source_label, want_redaction):
    file_status  = ""
    file_path_in = None

    if file_upload is not None:
        file_path_in = file_upload
        extracted, file_status = ingest_file(file_upload)
        if extracted:
            text = extracted

    if not text or not text.strip():
        empty_state = {"run_id": None, "entities": []}
        empty_msg = (
            '<div class="ds-banner ds-banner-warn">'
            'Please paste some text or upload a file to scan.</div>'
        )
        return (
            empty_msg, "", "",
            gr.update(visible=False, value=""),
            gr.update(visible=False, value=None),
            empty_state,
            gr.update(visible=False),
            gr.update(choices=[], value=[]),
            "",
        )

    if len(text) > CHAR_LIMIT:
        text = text[:CHAR_LIMIT]
        file_status = (file_status + f"  ⚠ Truncated to {CHAR_LIMIT:,} chars").strip()

    result = detector.detect(text, source_label=source_label or "gradio_ui")

    summary    = render_summary_html(result, file_status)
    entities   = render_entities_html(result.entities)
    annotated  = render_annotated_html(text, result.entities)

    # ── Redacted output — smart routing based on input type + user opt-in ─────
    redacted_text_update = gr.update(visible=False, value="")
    redacted_file_update = gr.update(visible=False, value=None)

    if want_redaction and result.entities:
        if file_path_in:
            redacted_path = build_redacted_file(file_path_in, text, result.entities)
            redacted_file_update = gr.update(visible=True, value=redacted_path)
        else:
            redacted_text_update = gr.update(
                visible=True,
                value=build_redacted_text(text, result.entities)
            )
    elif want_redaction and not result.entities:
        redacted_text_update = gr.update(
            visible=True,
            value="✓ Nothing to redact — no PII/PHI detected."
        )

    # ── Feedback state ────────────────────────────────────────────────────────
    entity_choices = [_entity_label(e) for e in result.entities]
    run_state = {
        "run_id": result.run_id,
        "entities": [
            {
                "entity_id":       e.entity_id,
                "entity_text":     e.text,
                "entity_type":     e.entity_type,
                "detection_layer": e.detection_layer,
                "confidence":      e.confidence,
                "label":           _entity_label(e),
            }
            for e in result.entities
        ],
    }

    feedback_visible = len(result.entities) > 0

    return (
        summary,
        entities,
        annotated,
        redacted_text_update,
        redacted_file_update,
        run_state,
        gr.update(visible=feedback_visible),
        gr.update(choices=entity_choices, value=[]),
        "",
    )


# ── Feedback handler ──────────────────────────────────────────────────────────

def submit_feedback(flagged_labels, run_state):
    if not run_state or not run_state.get("run_id"):
        return "⚠ No active scan to submit feedback for."
    if not flagged_labels:
        return "ℹ No entities selected — tick any incorrectly detected entities above, then submit."

    flagged_set = set(flagged_labels)
    to_log = []
    for ent in run_state["entities"]:
        if ent["label"] in flagged_set:
            to_log.append({
                "entity_id":         ent["entity_id"],
                "entity_text":       ent["entity_text"],
                "entity_type":       ent["entity_type"],
                "detection_layer":   ent["detection_layer"],
                "confidence":        ent["confidence"],
                "is_false_positive": True,
            })

    count = detector.audit.log_feedback(run_state["run_id"], to_log)
    if count:
        types = ", ".join(e["entity_type"] for e in to_log)
        return f"✓ Logged {count} false positive(s): **{types}**. Thank you — this helps improve accuracy."
    return "⚠ Feedback could not be saved."


def load_sample(sample_name):
    return SAMPLES.get(sample_name, "")


# ── UI ─────────────────────────────────────────────────────────────────────────

CUSTOM_CSS = """
/* DataSentry — theme-adaptive custom polish.
   Uses Gradio CSS variables so colors follow the active light/dark theme. */

.ds-hero {
  background: linear-gradient(135deg, #4f46e5 0%, #7c3aed 100%);
  color: #ffffff;
  padding: 28px 32px;
  border-radius: 14px;
  margin-bottom: 18px;
  box-shadow: 0 4px 16px rgba(79, 70, 229, 0.25);
}
.ds-hero h1 { margin: 0 0 6px 0; font-size: 28px; font-weight: 700; letter-spacing: -0.5px; color: #ffffff !important; }
.ds-hero p  { margin: 0; opacity: 0.92; font-size: 15px; color: #ffffff !important; }
.ds-hero-badges { display: flex; gap: 8px; margin-top: 14px; flex-wrap: wrap; }
.ds-hero-badge {
  background: rgba(255,255,255,0.18);
  color: #ffffff !important;
  padding: 4px 10px;
  border-radius: 999px;
  font-size: 12px;
  font-weight: 500;
  backdrop-filter: blur(4px);
}

.ds-banner {
  background: var(--background-fill-secondary, #eff6ff);
  border-left: 3px solid #3b82f6;
  padding: 10px 14px;
  border-radius: 6px;
  font-size: 13px;
  color: var(--body-text-color, #1e3a8a);
  margin-bottom: 12px;
}
.ds-banner-warn { border-color: #ca8a04; }

.ds-tiles {
  display: grid;
  grid-template-columns: repeat(4, 1fr);
  gap: 12px;
  margin-bottom: 12px;
}
.ds-tile {
  background: var(--block-background-fill, #ffffff);
  border: 1px solid var(--border-color-primary, #e5e7eb);
  border-radius: 10px;
  padding: 14px 16px;
  text-align: center;
}
.ds-tile.pii { border-top: 3px solid #f59e0b; }
.ds-tile.phi { border-top: 3px solid #ef4444; }
.ds-tile-num { font-size: 26px; font-weight: 700; color: var(--body-text-color, #111827); line-height: 1; }
.ds-tile-num .ds-unit { font-size: 13px; font-weight: 500; color: var(--body-text-color-subdued, #6b7280); margin-left: 2px; }
.ds-tile-lbl { font-size: 12px; color: var(--body-text-color-subdued, #6b7280); margin-top: 6px; text-transform: uppercase; letter-spacing: 0.4px; }
.ds-meta { font-size: 12px; color: var(--body-text-color-subdued, #6b7280); margin-bottom: 16px; }
.ds-meta code { background: var(--background-fill-secondary, #f3f4f6); color: var(--body-text-color, inherit); padding: 1px 6px; border-radius: 3px; }

.ds-cards { display: grid; gap: 8px; }
.ds-card {
  background: var(--block-background-fill, #ffffff);
  border: 1px solid var(--border-color-primary, #e5e7eb);
  border-radius: 8px;
  padding: 10px 14px;
  border-left: 3px solid #9ca3af;
}
.ds-card-pii { border-left-color: #f59e0b; }
.ds-card-phi { border-left-color: #ef4444; }
.ds-card-head { display: flex; justify-content: space-between; align-items: center; margin-bottom: 4px; }
.ds-card-type { display: flex; gap: 8px; align-items: center; }
.ds-cat { font-size: 10px; font-weight: 700; padding: 2px 7px; border-radius: 4px; letter-spacing: 0.4px; }
.ds-cat-pii { background: #fef3c7; color: #92400e !important; }
.ds-cat-phi { background: #fee2e2; color: #991b1b !important; }
.ds-type { font-size: 12px; font-weight: 600; color: var(--body-text-color, #374151); font-family: ui-monospace, monospace; }
.ds-card-conf { display: flex; align-items: center; gap: 8px; }
.ds-conf-pct { font-size: 12px; font-weight: 700; min-width: 32px; text-align: right; }
.ds-conf-bar { width: 60px; height: 4px; background: var(--background-fill-secondary, #f3f4f6); border-radius: 2px; overflow: hidden; }
.ds-conf-bar-fill { height: 100%; border-radius: 2px; }
.ds-card-text { font-family: ui-monospace, monospace; font-size: 13px; color: var(--body-text-color, #111827); padding: 4px 0; word-break: break-all; }
.ds-card-foot { font-size: 11px; color: var(--body-text-color-subdued, #6b7280); display: flex; gap: 8px; align-items: center; }
.ds-badge { font-size: 10px; padding: 1px 6px; border-radius: 3px; font-weight: 600; }
.ds-badge-claude { background: #ede9fe; color: #5b21b6 !important; }

.ds-empty { text-align: center; padding: 32px 16px; background: var(--block-background-fill, #ffffff); border: 1px dashed var(--border-color-primary, #d1d5db); border-radius: 10px; }
.ds-empty-ico { font-size: 28px; color: #10b981; margin-bottom: 6px; }
.ds-empty-msg { font-size: 14px; font-weight: 600; color: var(--body-text-color, #111827); }
.ds-empty-sub { font-size: 12px; color: var(--body-text-color-subdued, #6b7280); margin-top: 2px; }

.ds-annot {
  background: var(--background-fill-secondary, #fafbfc);
  border: 1px solid var(--border-color-primary, #e5e7eb);
  border-radius: 8px;
  padding: 14px;
  font-size: 13px;
  line-height: 1.7;
  max-height: 360px;
  overflow-y: auto;
  color: var(--body-text-color, #111827);
}
.ds-annot pre { margin: 0; white-space: pre-wrap; word-wrap: break-word; font-family: ui-monospace, monospace; color: inherit; }
.ds-mark { padding: 2px 4px; border-radius: 3px; position: relative; cursor: help; color: #111827 !important; }
.ds-mark-pii { background: #fef3c7; border-bottom: 2px solid #f59e0b; }
.ds-mark-phi { background: #fee2e2; border-bottom: 2px solid #ef4444; }
.ds-mark-tag { font-size: 9px; font-weight: 700; opacity: 0.7; vertical-align: super; margin-left: 2px; letter-spacing: 0.3px; color: inherit; }

.ds-section-title { font-size: 13px; font-weight: 700; color: var(--body-text-color, #374151); text-transform: uppercase; letter-spacing: 0.6px; margin: 18px 0 10px 0; opacity: 0.8; }

.ds-footer { text-align: center; color: var(--body-text-color-subdued, #6b7280); font-size: 12px; padding: 24px 0 12px 0; border-top: 1px solid var(--border-color-primary, #e5e7eb); margin-top: 24px; }
.ds-footer a { color: #818cf8; text-decoration: none; font-weight: 500; }
.ds-footer a:hover { text-decoration: underline; }

/* Dark-mode specific boosts for hardcoded light-backdrop elements */
.dark .ds-cat-pii { background: #78350f; color: #fef3c7 !important; }
.dark .ds-cat-phi { background: #7f1d1d; color: #fee2e2 !important; }
.dark .ds-badge-claude { background: #4c1d95; color: #ede9fe !important; }
.dark .ds-mark-pii { background: rgba(245, 158, 11, 0.22); border-bottom-color: #fbbf24; }
.dark .ds-mark-phi { background: rgba(239, 68, 68, 0.22); border-bottom-color: #f87171; }
.dark .ds-mark { color: var(--body-text-color) !important; }
.dark .ds-banner { background: rgba(59, 130, 246, 0.12); }
"""

THEME = gr.themes.Soft(
    primary_hue="indigo",
    secondary_hue="slate",
    neutral_hue="slate",
)

with gr.Blocks(title="DataSentry — PII & PHI Detection", theme=THEME, css=CUSTOM_CSS) as demo:

    # ── Hero ──────────────────────────────────────────────────────────────────
    gr.HTML("""
<div class="ds-hero">
  <h1>🛡 DataSentry</h1>
  <p>Detect PII and PHI across documents in seconds — with full audit trail and human-in-the-loop feedback.</p>
  <div class="ds-hero-badges">
    <span class="ds-hero-badge">4-layer hybrid engine</span>
    <span class="ds-hero-badge">HIPAA / GDPR aligned</span>
    <span class="ds-hero-badge">PDF · XLSX · DOCX · CSV · TXT</span>
    <span class="ds-hero-badge">Claude-powered arbitration</span>
  </div>
</div>
""")

    run_state = gr.State({})

    with gr.Row(equal_height=False):

        # ── Left column — inputs ──────────────────────────────────────────────
        with gr.Column(scale=2, min_width=320):
            gr.HTML('<div class="ds-section-title">Input</div>')

            with gr.Row():
                sample_dropdown = gr.Dropdown(
                    choices=list(SAMPLES.keys()),
                    label="Try a sample",
                    value=None,
                    info="Pre-fills the textbox with a realistic example.",
                )

            text_input = gr.Textbox(
                label="Text to scan",
                placeholder="Paste any text containing potential PII or PHI here...",
                lines=8,
            )

            file_upload = gr.File(
                label="Or upload a document",
                file_types=[".pdf", ".csv", ".xlsx", ".xls", ".docx", ".doc", ".txt"],
                type="filepath",
            )

            with gr.Accordion("Advanced options", open=False):
                source_label = gr.Textbox(
                    label="Source label",
                    value="gradio_ui",
                    info="Tag this scan in the audit log (e.g. 'support_intake', 'contract_review').",
                )
                want_redaction = gr.Checkbox(
                    label="Generate redacted output after scan",
                    value=False,
                    info="Pasted text → inline copy. File upload → downloadable file.",
                )

            detect_btn = gr.Button("🔍  Run Detection", variant="primary", size="lg")

        # ── Right column — outputs ────────────────────────────────────────────
        with gr.Column(scale=3, min_width=400):
            gr.HTML('<div class="ds-section-title">Results</div>')

            summary_out = gr.HTML(value=(
                '<div class="ds-empty">'
                '<div class="ds-empty-ico">⤴</div>'
                '<div class="ds-empty-msg">Ready when you are</div>'
                '<div class="ds-empty-sub">Pick a sample, paste text, or upload a document — then hit Run Detection.</div>'
                '</div>'
            ))
            entities_out = gr.HTML()

            with gr.Accordion("📍 Annotated source text", open=False):
                annotated_out = gr.HTML()

            redacted_text_out = gr.Textbox(
                label="Redacted text  (copy-safe)",
                interactive=False,
                lines=6,
                show_copy_button=True,
                visible=False,
            )
            redacted_file_out = gr.File(
                label="Download redacted file",
                visible=False,
            )

            # ── Feedback section — hidden until entities exist ────────────────
            with gr.Column(visible=False) as feedback_section:
                gr.HTML('<div class="ds-section-title">Help us improve</div>')
                gr.Markdown(
                    "Did we get anything wrong? Tick any **incorrectly detected** entities and submit. "
                    "Your feedback is logged to the audit trail and used to improve future accuracy."
                )
                fp_checkboxes = gr.CheckboxGroup(
                    choices=[],
                    label="Detected entities — tick false positives",
                    value=[],
                )
                feedback_btn    = gr.Button("Submit feedback", variant="secondary", size="sm")
                feedback_status = gr.Markdown("")

    # ── Architecture reference ────────────────────────────────────────────────
    with gr.Accordion("⚙️  How it works (4-layer architecture)", open=False):
        gr.Markdown("""
| Layer | Method | Confidence | Handles |
|-------|--------|------------|---------|
| **1. Regex** | 30+ keyword-gated patterns | 0.78–0.95 | SSN, email, credit card, ICD codes, international IDs (UK NHS, Indian Aadhaar, EU VAT, …) |
| **2. spaCy NER** | `en_core_web_sm` | 0.50–0.70 | Person names, locations, organizations, facilities |
| **3. Claude LLM** | Arbitration on ambiguous entities | < 0.75 threshold | Refines types, drops false positives, contextualizes medical terms |
| **4. SQLite audit** | Full provenance log | — | Every detection + every user feedback flag, queryable |

**Why this stack?** Regex catches structured identifiers cheaply. spaCy adds free-text recall for names and places. Claude handles the long-tail of ambiguous cases where context decides (e.g., "metformin" in a patient note vs a research blog). SQLite makes the whole pipeline auditable.
""")

    gr.HTML("""
<div class="ds-footer">
  Built by <a href="https://linkedin.com/in/srijan24" target="_blank">Srijan Gupta</a> ·
  <a href="https://github.com/sri-25/datasentry" target="_blank">GitHub</a> ·
  Powered by Claude · spaCy · Gradio
</div>
""")

    # ── Event wiring ──────────────────────────────────────────────────────────

    sample_dropdown.change(
        fn=load_sample,
        inputs=sample_dropdown,
        outputs=text_input,
    )

    detect_btn.click(
        fn=detect,
        inputs=[text_input, file_upload, source_label, want_redaction],
        outputs=[
            summary_out,
            entities_out,
            annotated_out,
            redacted_text_out,
            redacted_file_out,
            run_state,
            feedback_section,
            fp_checkboxes,
            feedback_status,
        ],
    )

    feedback_btn.click(
        fn=submit_feedback,
        inputs=[fp_checkboxes, run_state],
        outputs=[feedback_status],
    )


if __name__ == "__main__":
    demo.launch()
